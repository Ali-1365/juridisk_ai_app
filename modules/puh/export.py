import os
import sqlite3
from docx import Document
from fpdf import FPDF
from datetime import datetime

def exportera_till_docx(pm_text, filnamn="rättsfallsanalys.docx"):
    doc = Document()
    doc.add_heading("Rättsfallsanalys – Juridiskt AI-system", level=1)
    doc.add_paragraph(pm_text)
    doc.save(filnamn)
    print(f"[✔] DOCX skapad: {filnamn}")

def exportera_till_pdf(pm_text, filnamn="rättsfallsanalys.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in pm_text.split("\n"):
        pdf.cell(200, 10, txt=line.strip(), ln=True)
    pdf.output(filnamn)
    print(f"[✔] PDF skapad: {filnamn}")

def spara_till_sqlite(publicering_data, analys_text, db_path="rattsfall.db"):
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS rattsfall (
            id TEXT PRIMARY KEY,
            malnummer TEXT,
            domstol TEXT,
            datum TEXT,
            titel TEXT,
            analys TEXT
        )
    ''')
    data = (
        publicering_data.get("id"),
        publicering_data.get("malnummer"),
        publicering_data.get("domstol"),
        publicering_data.get("datum"),
        publicering_data.get("titel"),
        analys_text
    )
    c.execute("REPLACE INTO rattsfall VALUES (?, ?, ?, ?, ?, ?)", data)
    conn.commit()
    conn.close()
    print(f"[✔] Rättsfall sparat i SQLite: {db_path}")
